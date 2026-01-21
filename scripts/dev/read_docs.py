#!/usr/bin/env python3
import docx
import sys

def read_docx(filename):
    try:
        doc = docx.Document(filename)
        print(f'\n=== {filename} ===\n')
        
        # Read paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        
        # Also check tables
        if doc.tables:
            print('\n--- Tables ---')
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text for cell in row.cells])
                    if row_text.strip():
                        print(row_text)
    except Exception as e:
        print(f'Error reading {filename}: {e}')

# Read both files
read_docx('填写范本提供.docx')
read_docx('填写范本提供 (1).docx')
